# -*- coding: utf-8 -*-
"""C2 正常测试：上传 → worker 解析 → 链式入库（切块+embedding+rag_chunk）→ 状态/块数/向量断言。
前置：网关(8082) + 新版 parse worker 已启动。用法: python scripts/debug/verify_rag_c2_normal.py
"""
import io
import sys
import tempfile
import time
import uuid
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parents[2] / "rag-python" / "src"))

import psycopg
import psycopg.rows
import requests
from docx import Document as WDoc
from pgvector.psycopg import register_vector

import config

BASE = "http://localhost:8082"
fail = 0



def check(name, cond, detail=""):
    global fail
    if cond:
        print(f"[OK] {name}")
    else:
        print(f"[FAIL] {name} -> {detail}")
        fail += 1


def db():
    conn = psycopg.connect(config.PG_DSN, row_factory=psycopg.rows.dict_row)
    register_vector(conn)
    return conn


def new_user():
    u = f"c2n_{uuid.uuid4().hex[:10]}"
    h = {"X-Forwarded-For": "198.51.100.77"}
    requests.post(f"{BASE}/api/auth/register", json={"username": u, "password": "Passw0rd1"},
                  headers=h, timeout=10)
    r = requests.post(f"{BASE}/api/auth/login", json={"username": u, "password": "Passw0rd1"},
                      headers=h, timeout=10).json()
    return {"user": u, "token": r["token"], "headers": {"Authorization": f"Bearer {r['token']}"}}


def upload(u, path, filename):
    with open(path, "rb") as f:
        r = requests.post(f"{BASE}/api/files/upload",
                          files={"file": (filename, f)}, headers=u["headers"], timeout=30)
    assert r.status_code == 200, f"upload failed: {r.status_code} {r.text}"
    return r.json()


def wait_task(file_id, timeout=90):
    """轮询 parse_tasks 至终态，返回任务行。"""
    deadline = time.time() + timeout
    with db() as conn:
        while time.time() < deadline:
            row = conn.execute(
                "SELECT status, chunk_count, node_count, error FROM parse_tasks WHERE file_id=%s",
                (file_id,)).fetchone()
            if row and row["status"] not in ("pending", "parsing"):
                return row
            time.sleep(2)
    return None


tmp = Path(tempfile.mkdtemp(prefix="rag_c2_n_"))

# 1. txt（md 格式：标题 + 长段 + 列表）
p = tmp / "产品说明.md"
p.write_text(
    "# 产品说明\n\n## 2.1 检索原理\n\n本系统采用混合检索架构，融合向量语义检索与关键词精确匹配两路召回，"
    "再通过倒数排名融合算法合并排序结果。向量检索擅长语义近义匹配，关键词检索擅长专有名词匹配，"
    "两者互补可显著提升召回准确率。\n\n"
    "- 召回阶段：向量 Top50 与关键词 Top50 并行\n"
    "- 融合阶段：RRF 倒数排名融合\n"
    "- 精排阶段：交叉编码器重排序\n",
    encoding="utf-8")

# 2. docx（标题 + 大表格）
d = tmp / "规范.docx"
wd = WDoc()
wd.add_heading("技术规范", level=1)
wd.add_heading("3.2 参数要求", level=2)
wd.add_paragraph("以下为各组件参数要求，涉及内存、磁盘与并发三项指标。")
tbl = wd.add_table(rows=14, cols=3)
tbl.rows[0].cells[0].text = "组件"
tbl.rows[0].cells[1].text = "内存"
tbl.rows[0].cells[2].text = "并发"
for i in range(1, 14):
    tbl.rows[i].cells[0].text = f"组件{i}"
    tbl.rows[i].cells[1].text = f"{i}GB"
    tbl.rows[i].cells[2].text = f"{i * 100}TPS"
wd.save(d)

u = new_user()
f1 = upload(u, p, "产品说明.md")
f2 = upload(u, d, "规范.docx")
check("上传返回 id", f1["id"] > 0 and f2["id"] > 0, (f1, f2))

t1 = wait_task(f1["id"])
t2 = wait_task(f2["id"])
check("txt 状态 success", t1 and t1["status"] == "success", t1)
check("docx 状态 success", t2 and t2["status"] == "success", t2)

# 3. rag_chunk 断言
with db() as conn:
    def chunk_rows(file_id):
        return conn.execute(
            "SELECT chunk_type, chars, heading_path, embedding, embed_model "
            "FROM rag_chunk WHERE file_id=%s ORDER BY seq", (file_id,)).fetchall()

    rows1 = chunk_rows(f1["id"])
    rows2 = chunk_rows(f2["id"])

check("txt 块数=chunk_count", len(rows1) == (t1["chunk_count"] or 0),
      f"{len(rows1)} vs {t1['chunk_count']}")
check("docx 块数=chunk_count", len(rows2) == (t2["chunk_count"] or 0),
      f"{len(rows2)} vs {t2['chunk_count']}")
check("txt 块数>0", len(rows1) > 0)
check("docx 块数>0", len(rows2) > 0)
check("txt 含列表块", any(r["chunk_type"] == "list" for r in rows1))
check("docx 含表格块", any(r["chunk_type"] == "table" for r in rows2))

# 4. embedding 契约：非空、768 维、模型名
vec1 = rows1[0]["embedding"]
check("embedding 非空", vec1 is not None)
check("embedding 768 维", vec1 is not None and len(vec1) == 768,
      None if vec1 is None else len(vec1))
check("embed_model 记录", all(r["embed_model"] == "bge-base-zh-v1.5-onnx-int8" for r in rows1),
      {r["embed_model"] for r in rows1})

# 5. 标题路径注入（docx 的块带标题路径）
paths = {r["heading_path"] for r in rows2}
check("docx 标题路径", any(p.startswith("技术规范 > 3.2 参数要求") for p in paths), paths)

# 6. 幂等：同文件重复 ingest 块数不变（ON CONFLICT + 删旧写新）
import json as _json

from ingest.indexer import ingest as do_ingest
from ingest.parser.base import DocumentNode

payload = _json.loads((config.PARSED_DIR / f"{f1['id']}.json").read_text(encoding="utf-8"))
nodes = [DocumentNode(n["type"], n["text"], n["meta"]) for n in payload["nodes"]]
cnt1 = do_ingest(f1["id"], nodes)
cnt2 = do_ingest(f1["id"], nodes)
check("重复入库块数不变", cnt1 == cnt2 == len(rows1), (cnt1, cnt2, len(rows1)))

print(f"\nC2 normal: {'PASS' if fail == 0 else f'{fail} FAILED'}")
sys.exit(1 if fail else 0)

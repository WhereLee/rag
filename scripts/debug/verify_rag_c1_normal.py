# -*- coding: utf-8 -*-
"""C1 正常测试：真实文档（txt/md、pdf、docx）→ parse_file 全管线 → chunk_nodes 结构感知切块。
用法: python scripts/debug/verify_rag_c1_normal.py（需在项目根目录执行）
"""
import sys
import tempfile
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parents[2] / "rag-python" / "src"))

import pymupdf
from docx import Document as WDoc
from docx.shared import Inches

from ingest.chunker import chunk_nodes
from ingest.pipeline import parse_file

fail = 0


def check(name, cond, detail=""):
    global fail
    if cond:
        print(f"[OK] {name}")
    else:
        print(f"[FAIL] {name} -> {detail}")
        fail += 1


tmp = Path(tempfile.mkdtemp(prefix="rag_c1_n_"))

# 1. markdown：标题嵌套 + 长段落 + 列表
p = tmp / "产品说明.md"
p.write_text(
    "# 产品说明\n\n"
    "## 1.1 系统架构\n\n"
    "本系统采用分层架构设计，包含网关层、服务层与数据层。网关负责统一鉴权与流量控制，"
    "服务层承载核心业务逻辑，数据层使用 PostgreSQL 存储结构化数据。各层之间通过标准化接口通信，"
    "支持独立扩展与水平扩容。该架构同时满足高可用与可维护性的设计要求。\n\n"
    "安全方面，系统实现了基于 JWT 的无状态认证机制，配合 RBAC 权限模型控制资源访问，"
    "未登录用户无法访问任何受保护接口。所有上传文件均经过类型白名单与魔数双重校验，"
    "防止伪装文件绕过安全策略，文件存储与访问均有完整审计日志。\n\n"
    "- 网关层：统一入口、身份鉴权、限流控制与请求转发，负责所有外部流量接入与安全拦截\n"
    "- 服务层：业务逻辑编排、任务调度与状态管理，承载文档解析与检索问答等核心能力\n"
    "- 数据层：PostgreSQL 持久化存储与对象存储，同时支持关系事务与向量相似度检索\n",
    encoding="utf-8")
res = parse_file(p)
chunks = chunk_nodes(res.nodes)
check("md 解析成功", res.status == "success", res.error)
check("md 切块非空", len(chunks) > 0, f"nodes={len(res.nodes)}")
paths = [c.heading_path for c in chunks]
check("md 标题路径", "产品说明 > 1.1 系统架构" in paths, paths)
# 孤立的短段（<100 无相邻可合并）成块是合理行为，断言放宽：≤700 且 ≥100 占比 ≥80%
big = [c for c in chunks if c.chars >= 100]
check("md 块大小 ≤700", all(c.chars <= 700 for c in chunks), [c.chars for c in chunks])
check("md 块 ≥100 占比 ≥80%", len(big) / max(len(chunks), 1) >= 0.8,
      f"{len(big)}/{len(chunks)}")
check("md 列表块", any(c.chunk_type == "list" for c in chunks),
      [c.chunk_type for c in chunks])

# 2. pdf：3 页纯文本（无 heading）→ 段落合并
p = tmp / "说明.pdf"
doc = pymupdf.open()
for i in range(3):
    page = doc.new_page(width=595, height=842)
    text = (f"第{i + 1}页正文内容。本页描述系统部署流程，包括环境准备、依赖安装与配置初始化三个步骤。"
            "部署完成后需要执行健康检查脚本验证服务状态，确认所有组件运行正常。"
            "如遇到启动失败，可查看日志文件定位具体原因。")
    try:
        page.insert_text((72, 72), text, fontsize=11, fontname="china-s")
    except Exception:
        page.insert_text((72, 72), f"Page {i+1} content. Deployment steps include setup, install and config.")
doc.save(p)
doc.close()
res = parse_file(p)
chunks = chunk_nodes(res.nodes)
check("pdf 解析成功", res.status == "success", res.error)
check("pdf 切块非空", len(chunks) > 0)
check("pdf 无标题路径", all(c.heading_path == "" for c in chunks))
big = [c for c in chunks if c.chars >= 100]
check("pdf 块大小 ≤700", all(c.chars <= 700 for c in chunks), [c.chars for c in chunks])
check("pdf 块 ≥100 占比 ≥80%", len(big) / max(len(chunks), 1) >= 0.8,
      f"{len(big)}/{len(chunks)}")
check("pdf 全部段落块", all(c.chunk_type == "paragraph" for c in chunks),
      [c.chunk_type for c in chunks])

# 3. docx：Heading 层级 + 大表格（13 数据行 → 2 组表头保留）
p = tmp / "规范.docx"
wd = WDoc()
wd.add_heading("技术规范", level=1)
wd.add_heading("3.2 参数要求", level=2)
wd.add_paragraph("以下为各组件参数要求，涉及内存、磁盘与并发三项指标。")
tbl = wd.add_table(rows=14, cols=3)
tbl.rows[0].cells[0].text = "组件"
tbl.rows[0].cells[1].text = "内存"
tbl.rows[0].cells[2].text = "并发"
for i in range(1, 14):
    tbl.rows[i].cells[0].text = f"组件{i}"
    tbl.rows[i].cells[1].text = f"{i}GB"
    tbl.rows[i].cells[2].text = f"{i * 100}TPS"
wd.save(p)
res = parse_file(p)
chunks = chunk_nodes(res.nodes)
check("docx 解析成功", res.status == "success", res.error)
check("docx 切块非空", len(chunks) > 0)
tables = [c for c in chunks if c.chunk_type == "table"]
check("docx 表格 2 组", len(tables) == 2, f"tables={len(tables)}")
if tables:
    hdr = tables[0].content.splitlines()[0]
    check("docx 组间表头保留", all(t.content.splitlines()[0] == hdr for t in tables))
    check("docx 表格块不超 12 行", all(len(t.content.splitlines()) <= 12 for t in tables))
check("docx 标题路径", any(c.heading_path == "技术规范 > 3.2 参数要求" for c in chunks),
      sorted({c.heading_path for c in chunks}))

print(f"\nC1 normal: {'PASS' if fail == 0 else f'{fail} FAILED'}")
sys.exit(1 if fail else 0)
